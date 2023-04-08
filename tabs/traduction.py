from tempfile import NamedTemporaryFile

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Cm

from .utils import *

import re
import streamlit_select_image as stsi
import docx


def traduction():
    if "table" not in st.session_state:
        load_table()
    if "traduction_sentences" not in st.session_state:
        st.session_state["traduction_sentences"] = pd.DataFrame(
            {"sentence": [""], "words": [[]], "images": [[]], "columns": [{}], "before": [""]})
    if "traduction_count" not in st.session_state:
        st.session_state["traduction_count"] = 1
    st.markdown("#### Titre du document")
    title = st.text_input("Titre du document :",
                          label_visibility="collapsed",
                          key="traduction_title")
    st.markdown("#### Phrases")
    generate_sentences()
    generate_document(title=title)


# UTILS

def add_sentence():
    st.session_state["traduction_sentences"].loc[st.session_state["traduction_count"]] = "", [], [], {}, False
    st.session_state["traduction_count"] += 1


def delete_sentence(_id: str):
    def delete_sentence_id():
        st.session_state["traduction_sentences"].drop(index=_id, inplace=True)
        if f"traduction_remove_{_id}" in st.session_state:
            del st.session_state[f"traduction_remove_{_id}"]
        if f"traduction_sentence_{_id}" in st.session_state:
            del st.session_state[f"traduction_sentence_{_id}"]
        stsi_keys = [key for key in st.session_state if re.match("^traduction_sentence_[0-9]+_[0-9]+_[0-9]+_.+", key)]
        for key in stsi_keys:
            if key in st.session_state:
                del st.session_state[key]

    return delete_sentence_id


def generate_document(title):
    stream = None
    nb_line = 1
    buttons_columns_trad = st.columns(2)
    with buttons_columns_trad[0]:
        if st.button("Valider"):
            with st.spinner("Génération du document Word..."):
                document_trad = docx.Document("static/template.docx")
                # Title generation
                for p in document_trad.paragraphs:
                    if p.style.name == "Title":
                        p.text = title
                        p.style.font.color.rgb = RGBColor(0, 0, 0)
                # Sentences generation
                for _id in st.session_state["traduction_sentences"].index:
                    if st.session_state["traduction_sentences"].at[_id, "sentence"] != "":
                        document_trad.add_paragraph(st.session_state["traduction_sentences"].at[_id, "sentence"])
                        nb_line += 1
                        for i in range(0, len(st.session_state["traduction_sentences"].at[_id, "words"]), 8):
                            j = i
                            grid = document_trad.add_paragraph()
                            r = grid.add_run()
                            while j < len(st.session_state["traduction_sentences"].at[_id, "words"]) and j < i + 8:
                                words = st.session_state["traduction_sentences"].at[_id, "words"][j]
                                search_si = f"traduction_sentence_{_id}_{i}_{j % 8}_{words}"
                                if search_si in st.session_state:
                                    if st.session_state[search_si] != "":
                                        r.add_picture(f"{st.session_state[search_si]}", width=Cm(3.05), height=Cm(3.05))
                                j += 1

                section = document_trad.sections[-1]
                footer = section.footer
                p = footer.paragraphs[0]
                p.text = "Fait grâce à "
                r = p.add_run()
                r.add_picture("static/logo.png", width=Cm(1.5), height=Cm(0.96))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                with NamedTemporaryFile() as tmp:
                    document_trad.save(tmp.name)
                    tmp.seek(0)
                    stream = tmp.read()

    if stream:
        with buttons_columns_trad[1]:
            download_trad = st.download_button(label="Télécharger", data=stream,
                                               file_name=f"pictortho_{tmp.name.split('.')[0].split('/')[-1]}.docx",
                                               mime="application/msword", key="dl_trad")


def generate_sentences():
    list_keys = list(st.session_state["traduction_sentences"].index)
    for i in range(len(list_keys)):
        _id = list_keys[i]
        line = st.columns([33, 1, 1])
        with line[0]:
            st.session_state["traduction_sentences"].at[_id, "sentence"] = st.text_input(
                label=f"traduction_sentence_{_id}", value="",
                label_visibility="collapsed",
                key=f"traduction_sentence_{_id}")
        with line[1]:
            if len(list_keys) != 1:
                st.button(label=" \- ", key=f"traduction_remove_{_id}", on_click=delete_sentence(_id))
        if i == len(list_keys) - 1:
            with line[2]:
                st.button(label="\+", key="add_sentence", on_click=lambda: add_sentence())
        if st.session_state["traduction_sentences"].at[_id, "sentence"] != "":
            if st.session_state["traduction_sentences"].at[_id, "before"] != st.session_state["traduction_sentences"].at[_id, "sentence"]:
                with st.spinner("Traduction en cours..."):
                    res = translate(st.session_state["traduction_sentences"].at[_id, "sentence"])
                    st.session_state["traduction_sentences"].at[_id, "words"] = [r["words"] for r in res]
                    st.session_state["traduction_sentences"].at[_id, "images"] = [list(r["data"]["image"].unique()) for r in
                                                                                  res]
                st.session_state["traduction_sentences"].at[_id, "before"] = st.session_state["traduction_sentences"].at[_id, "sentence"]
            for j in range(0, len(st.session_state["traduction_sentences"].at[_id, "words"]), 8):
                st.session_state["traduction_sentences"].at[_id, "columns"][str(j)] = st.columns(8)
                for k in range(len(st.session_state["traduction_sentences"].at[_id, "words"][j:j + 8])):
                    name = st.session_state["traduction_sentences"].at[_id, "words"][j:j + 8][k]
                    images = st.session_state["traduction_sentences"].at[_id, "images"][j:j + 8][k]
                    with st.session_state["traduction_sentences"].at[_id, "columns"][str(j % 8)][k % 8]:
                        select_img = stsi.st_select_image(
                            options=images,
                            label=name, no_choice=True,
                            key=f"traduction_sentence_{_id}_{j}_{k}_{name}")
