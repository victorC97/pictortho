from .utils import *
import streamlit_select_image as stsi
import docx
from tempfile import NamedTemporaryFile
from io import BytesIO
from docx.shared import Cm
from PIL import Image
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

WIDTH = 4.41
HEIGHT = 4.80


def tla():
    st.markdown("#### Configuration")
    parameters = st.columns(6)
    with parameters[0]:
        st.write("Nombre de colonnes :")
    with parameters[1]:
        nb_col = st.number_input(label="Nombre de colonnes :", min_value=1, value=5, max_value=6,
                                 label_visibility="collapsed", key="nb_col")
    with parameters[2]:
        st.write("Nombre de lignes :")
    with parameters[3]:
        nb_row = st.number_input(label="Nombre de lignes :", min_value=1, value=4, max_value=6,
                                 label_visibility="collapsed", key="nb_row")
    with parameters[5]:
        reset_btn = st.button(label="Réinitialiser", on_click=lambda: reset_tla(), key="tla_reset")
    columns = st.columns(nb_col)
    for i in range(0, nb_col):
        for j in range(0, nb_row):
            with columns[i]:
                text = st.text_input(label=f"tla_text_{i}_{j}", label_visibility="hidden", key=f"tla_text_{i}_{j}")
    stream = None

    st.markdown("#### Tableau de Langage Assisté")
    tableau = st.columns(nb_col)
    for i in range(0, nb_col):
        for j in range(0, nb_row):
            with tableau[i]:
                res = translate(st.session_state[f"tla_text_{i}_{j}"])
                if len(res) == 0:
                    pass
                elif len(res) == 1:
                    choice = st.selectbox(options=[res[0]["words"]], label=f"tla_choice_{i}_{j}",
                                          label_visibility="collapsed", key=f"tla_choice_{i}_{j}", disabled=True)
                    tla_select_img = stsi.st_select_image(
                        options=res[0]["data"]["image"].to_list(),
                        label=f"{res[0]['words']}", no_choice=True,
                        key=f"tla_img_{i}_{j}_{res[0]['words']}")
                else:
                    choice = st.selectbox(options=[el["words"] for el in res], label=f"tla_choice_{i}_{j}",
                                          label_visibility="collapsed", key=f"tla_choice_{i}_{j}")
                    infos = [el for el in res if el["words"] == choice][0]
                    tla_select_img = stsi.st_select_image(
                        options=infos["data"]["image"].to_list(),
                        label=f"{infos['words']}", no_choice=True,
                        key=f"tla_img_{i}_{j}_{infos['words']}")

    document_tla = None
    button_columns_tla = st.columns(2)
    if st.button("Valider", key="tla_generate"):
        with st.spinner("Loading..."):
            document_tla = docx.Document()
            tla = document_tla.add_table(rows=nb_col, cols=nb_row)
            tla.style = 'Table Grid'
            tla.autofit = False
            tla.allow_autofit = False
            for i in range(0, nb_col):
                for j in range(0, nb_row):
                    cell = tla.cell(i, nb_row - j - 1)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cell.width = Cm(get_width(nb_row))
                    p = cell.paragraphs[0]
                    p.add_run(st.session_state[f"tla_text_{i}_{j}"] + "\n\n").bold = True
                    cell.add_paragraph('')
                    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                    r = p.add_run("")
                    patrn = f"tla_img_{i}_{j}_"
                    pkey = [key for key in st.session_state if re.match(patrn, key)]
                    if len(pkey) > 0:
                        if st.session_state[f"tla_text_{i}_{j}"] != "" and st.session_state[pkey[0]] != "":
                            io_img = BytesIO()
                            img = Image.open(st.session_state[pkey[0]]).rotate(270)
                            img.save(io_img, "PNG")
                            r.add_picture(io_img, width=Cm(0.7 * get_width(nb_row)), height=Cm(0.7 * get_width(nb_row)))
                    sz = 36
                    set_cell_border(
                        cell,
                        top={"sz": sz, "val": "single", "color": "#000000", "space": "0"},
                        bottom={"sz": sz, "val": "single", "color": "#000000"},
                        start={"sz": sz, "val": "single", "color": "#000000"},
                        end={"sz": sz, "val": "single", "color": "#000000", },
                    )
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    textDirection = OxmlElement('w:textDirection')
                    textDirection.set(qn('w:val'), 'tbRl')  # btLr tbRl
                    tcPr.append(textDirection)

            for row in tla.rows:
                row.height = Cm(get_height(nb_col))

            sections = document_tla.sections
            margin = 1.27
            for section in sections:
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)

            with NamedTemporaryFile() as tmp:
                document_tla.save(tmp.name)
                tmp.seek(0)
                stream = tmp.read()

    with button_columns_tla[1]:
        if stream:
            download_tla = st.download_button(label="Télécharger", data=stream,
                                              file_name=f"pictortho_{tmp.name.split('.')[0].split('/')[-1]}.docx",
                                              mime="application/msword", key="tla_dl")


def reset_tla():
    to_remove = [key for key in st.session_state if re.match("tla_", key)]
    to_reset = [key for key in st.session_state if re.match("tla_text_", key)]
    for key in to_remove:
        del st.session_state[key]
    for key in to_reset:
        st.session_state[key] = ""


def get_width(nb_row):
    return 19 / nb_row


def get_height(nb_col):
    return 25.7 / nb_col
