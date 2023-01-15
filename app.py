from io import BytesIO
import streamlit as st
import pandas as pd
import re
import docx
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tempfile import NamedTemporaryFile
import copy
import base64
import streamlit_select_image as stsi
from streamlit_elements import elements, mui, html, lazy, sync
from st_aggrid import AgGrid
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

# HEADERS AND BEAUTIFY
st.set_page_config(page_title="Pictortho", layout="wide", page_icon="icon.png")

with open("static/style.css") as css:
    st.markdown(f'<style>{css.read()}</style>', unsafe_allow_html=True)

with open('static/logo.svg', 'r') as svg:
    b64 = base64.b64encode(svg.read().encode('utf-8')).decode("utf-8")

with open('static/logo_ARASAAC.png', 'rb') as f:
    arasaac = base64.b64encode(f.read())

with open("README.md", "r") as f:
    readme = f.readlines()

# NEEDED VARIABLES

if "table" not in st.session_state:
    st.session_state.table = pd.read_csv("static/table.csv", sep="\t", index_col=0)
    # GETTING LISTS
    commas_columns = ["_id", "Pictogramme", "Conjugaison", "schematic", "sex", "violence", "aac", "aacColor", "skin",
                      "hair", "downloads"]
    double_point_columns = ["Catégories"]
    for column in commas_columns:
        st.session_state.table[column] = st.session_state.table[column].apply(
            lambda string: string.split(",") if type(string) is str else string)
    for column in double_point_columns:
        st.session_state.table[column] = st.session_state.table[column].apply(
            lambda string: string.split(":") if type(string) is str else string)

if "verbs" not in st.session_state:
    st.session_state.verbs = pd.read_csv("static/conjugaison.csv", sep=";").set_index("Mots conjugués")
    st.session_state.verbs = st.session_state.verbs[~st.session_state.verbs.index.duplicated(keep='first')]

if "sentences" not in st.session_state:
    st.session_state.sentences = {"0": {"id": 0, "role": "add"}}

if "count" not in st.session_state:
    st.session_state.count = 1

if "manuel_file" not in st.session_state:
    with open("static/manuel.md", "r") as f:
        st.session_state.manuel_file = f.read()

if "mode" not in st.session_state:
    st.session_state.mode = "pictortho"

if "dictionnaire_result" not in st.session_state:
    st.session_state.dictionnaire_result = st.session_state.table.iloc[0:20]


# UTILS
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_sentence():
    st.session_state.sentences[str(st.session_state.count)] = {"id": st.session_state.count, "role": "remove"}
    st.session_state.count += 1


def delete_sentence(i: int):
    def delete_sentencei():
        del st.session_state.sentences[str(i)]
        del st.session_state[f"sentence{i}"]
        del st.session_state[f"sentence{i}_remove"]

    return delete_sentencei


def get_pictos(sentence: str):
    res = []
    # Replace verb to standard form
    temp_sentence = copy.deepcopy(sentence)
    temp_sentence = temp_sentence.lower()
    temp_sentence = temp_sentence.replace(" l'", " l' ").replace("j'", "je ")
    temp_sentence = [el[:-1] if el[-1] == "." else el for el in re.split(" ", temp_sentence) if el != ""]
    # Match words
    res = {}
    j = 0
    while j < len(temp_sentence):
        for k in range(0, len(temp_sentence) - j):
            words = " ".join(temp_sentence[j:j + k + 1])
            if words == "l'":
                find_res = st.session_state.table[
                    st.session_state.table.index.str.fullmatch("le") | st.session_state.table.index.str.fullmatch("la")]
            else:
                find_res = st.session_state.table[st.session_state.table.index.str.fullmatch(words)]
            if find_res.shape[0] != 0:
                if str(j) not in res:
                    res[str(j)] = {"words": words, "data": find_res}
                else:
                    # Remove doublon aliments de dinette -> aliments de dinette + dinette : we need to remove dinette
                    if len(words) > len(res[str(j)]["words"]):
                        res[str(j)] = {"words": words, "data": find_res}
        j += len(res[str(j)]["words"].split(" ")) if str(j) in res else 1
    return [{"words": res[key]["words"], "data": res[key]["data"]} for key in res]


def set_mode(mode: str):
    st.session_state.mode = mode


def pictortho_page():
    st.markdown("#### Titre du document à générer:")

    title = st.text_input("Titre du document :", label_visibility="collapsed")

    st.markdown("#### Phrases à traduire :")

    stream = None
    nb_line = 1
    for key in st.session_state.sentences:
        info = st.session_state.sentences[key]
        line = st.columns([33, 1])
        with line[0]:
            st.session_state.sentences[key]["sentence"] = st.text_input(label=f"sentence{key}", value="",
                                                                        label_visibility="collapsed",
                                                                        key=f"sentence{key}")
        with line[1]:
            if info["role"] == "add":
                st.button(label="\+", key="add_sentence", on_click=lambda: add_sentence())
            elif info["role"] == "remove":
                st.button(label=" \- ", key=f"sentence{key}_remove", on_click=delete_sentence(key))
        st.session_state.sentences[key]["columns"] = []
        if st.session_state[f"sentence{key}"] != "":
            with st.spinner("Traduction en cours..."):
                st.session_state.sentences[key]["res"] = get_pictos(st.session_state[f"sentence{key}"])
            for i in range(0, len(st.session_state.sentences[key]["res"]), 8):
                st.session_state.sentences[key]["columns"].append(st.columns(8))
                for index in range(len(st.session_state.sentences[key]["res"][i:i + 8])):
                    with st.session_state.sentences[key]["columns"][i % 8][index % 8]:
                        infos = st.session_state.sentences[key]["res"][i:i + 8][index]
                        imgs = []
                        for l in infos["data"]["Pictogramme"].values:
                            for el in l:
                                if el not in imgs:
                                    imgs.append(el)
                        select_img = stsi.st_select_image(
                            options=[f"pictos/{img}" for img in imgs],
                            label=f"{infos['words']}", no_choice=True,
                            key=f"sentence{key}_si_{i}_{index}_{infos['words']}")
    buttons_columns_trad = st.columns(2)
    with buttons_columns_trad[0]:
        if st.button("Générer"):
            with st.spinner("Génération du document Word..."):
                document_trad = docx.Document("static/template.docx")
                for p in document_trad.paragraphs:
                    if p.style.name == "Title":
                        p.text = title
                        p.style.font.color.rgb = RGBColor(0, 0, 0)
                for key in st.session_state.sentences:
                    document_trad.add_paragraph(st.session_state[f"sentence{key}"])
                    nb_line += 1
                    grid = document_trad.add_paragraph()
                    r = grid.add_run()
                    for i in range(0, len(st.session_state.sentences[key]["res"]), 8):
                        for index in range(len(st.session_state.sentences[key]["res"][i:i + 8])):
                            search_si = f'sentence{key}_si_{i}_{index}'
                            sis = [key for key in st.session_state if re.search(search_si, key)]
                            if len(sis) == 1:
                                si = sis[0]
                            else:
                                st.error(sis)
                            if st.session_state[si] != "":
                                r.add_picture(f"{st.session_state[si]}", width=Cm(3.05), height=Cm(3.05))

                section = document_trad.sections[-1]
                footer = section.footer
                p = footer.paragraphs[0]
                p.text = "Fait grâce à "
                r = p.add_run()
                r.add_picture("static/logo.png", width=Cm(1.5), height=Cm(0.96))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                with NamedTemporaryFile() as tmp:
                    document_trad.save(tmp.name)
                    tmp.seek(0)
                    stream = tmp.read()

    if stream:
        with buttons_columns_trad[1]:
            download_trad = st.download_button(label="Télécharger", data=stream,
                                      file_name=f"pictortho_{tmp.name.split('.')[0].split('/')[-1]}.docx",
                                      mime="application/msword", key="dl_trad")


def tla_page():
    st.markdown("#### Configuration")
    parameters = st.columns(2)
    with parameters[0]:
        nb_col = st.number_input(label="Nombre de colonnes :", min_value=0, value=5, max_value=6, label_visibility="hidden", key="nb_col")
    with parameters[1]:
        nb_row = st.number_input(label="Nombre de lignes :", min_value=0, value=4, max_value=6, label_visibility="hidden", key="nb_row")

    columns = st.columns(nb_col)
    for i in range(0, nb_col):
        for j in range(0, nb_row):
            with columns[i]:
                text = st.text_input(label=f"grid_text_{i}_{j}", label_visibility="hidden", key=f"grid_text_{i}_{j}")

    stream = None

    st.markdown("#### Tableau de Langage Assisté")
    tableau = st.columns(nb_col)
    for i in range(0, nb_col):
        for j in range(0, nb_row):
            with tableau[i]:
                st.session_state[f"grid_res_{i}_{j}"] = get_pictos(st.session_state[f"grid_text_{i}_{j}"])
                for k in range(0, len(st.session_state[f"grid_res_{i}_{j}"])):
                    infos = st.session_state[f"grid_res_{i}_{j}"][k]
                    imgs = []
                    for l in infos["data"]["Pictogramme"].values:
                        for el in l:
                            if el not in imgs:
                                imgs.append(el)
                    tla_select_img = stsi.st_select_image(
                        options=[f"pictos/{img}" for img in imgs],
                        label=f"{infos['words']}", no_choice=True,
                        key=f"grid_res_{i}_{j}_{k}_{infos['words']}")
    document_tla = None
    button_columns_tla = st.columns(2)
    with button_columns_tla[0]:
        if st.button("Générer", key="generate_tla"):
            with st.spinner("Loading..."):
                document_tla = docx.Document()
                tla = document_tla.add_table(rows=nb_col, cols=nb_row)
                tla.style = 'Table Grid'
                tla.autofit = False
                tla.allow_autofit = False
                for i in range(0, nb_col):
                    for j in range(0, nb_row):
                        cell = tla.cell(i, nb_row - j - 1)
                        cell.width = Cm(4.41)
                        p = cell.paragraphs[0]
                        p.add_run(st.session_state[f"grid_text_{i}_{j}"] + "\n\n").bold = True
                        cell.add_paragraph('')
                        p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                        r = p.add_run("")
                        patrn = f"grid_res_{i}_{j}_"
                        res = [key for key in st.session_state if re.match(patrn, key)]
                        if st.session_state[res] != "":
                            io_img = BytesIO()
                            img = Image.open(st.session_state[res[0]]).rotate(270)
                            img.save(io_img, "PNG")
                            r.add_picture(io_img, width=Cm(2.5), height=Cm(2.5))
                            sz = 36
                            set_cell_border(
                                cell,
                                top={"sz": sz, "val": "single", "color": "#000000", "space": "0"},
                                bottom={"sz": sz, "val": "single", "color": "#000000"},
                                start={"sz": sz, "val": "single", "color": "#000000"},
                                end={"sz": sz, "val": "single", "color": "#000000",},
                            )
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        textDirection = OxmlElement('w:textDirection')
                        textDirection.set(qn('w:val'), 'tbRl')  # btLr tbRl
                        tcPr.append(textDirection)

                for row in tla.rows:
                    row.height = Cm(4.80)

                sections = document_tla.sections
                margin = 1.27
                for section in sections:
                    section.top_margin = Cm(margin)
                    section.bottom_margin = Cm(margin)
                    section.left_margin = Cm(margin)
                    section.right_margin = Cm(margin)

                with NamedTemporaryFile() as tmp:
                    document_tla.save(tmp.name)
                    tmp.seek(0)
                    stream = tmp.read()

    with button_columns_tla[1]:
        if stream:
            download_tla = st.download_button(label="Télécharger", data=stream,
                                      file_name=f"pictortho_{tmp.name.split('.')[0].split('/')[-1]}.docx",
                                      mime="application/msword", key="dl_tla")




def set_dictionnaire_result(text: str):
    st.session_state.dictionnaire_result = st.session_state.table[st.session_state.table.index.str.contains(text)]


def dictionnaire():
    st.markdown("#### Rechercher un pictogramme :")
    searching_text = st.text_input("Rechercher un pictogramme :", label_visibility="collapsed",
                                   key="search_dictionnaire")
    if searching_text != "":
        set_dictionnaire_result(searching_text)
    AgGrid(st.session_state.dictionnaire_result[["Mots écrits", "_id", "Catégories"]])


# APP START

with elements("app-bar"):
    with mui.Box(sx={"flexGrow": 1}, mx={0}):
        with mui.AppBar(position="static", sx={"backgroundColor": "#FFFFFF"}):
            with mui.Toolbar(display="flex", sx={"justifyContent": "space-between", "padding": 2}):
                mui.Box(
                    html.img(src=f"data:image/svg+xml;base64,{b64[:-2]}", alt="pictortho", width="120px"),
                    mui.Typography("écrit grâce à", variant="h4", color="black"),
                    html.img(src=f"data:image/png;base64,{arasaac.decode()}", width="300px"),
                    display="flex", alignItems="center"
                )
                mui.Box(
                    mui.Box(
                        mui.IconButton(
                            mui.icon.Create,
                            onClick=lambda: set_mode("pictortho"),
                            sx={"color": "black"}
                        ),
                        mui.Typography(
                            "Traduire",
                            variant="h6",
                            color="black"
                        ),
                        display="flex", alignItems="center", sx={"marginRight": "10px"}
                    ),
                    mui.Box(
                        mui.IconButton(
                            mui.icon.TableView,
                            onClick=lambda: set_mode("tla"),
                            sx={"color": "black"}
                        ),
                        mui.Typography(
                            "TLA",
                            variant="h6",
                            color="black"
                        ),
                        display="flex", alignItems="center", sx={"marginRight": "10px"}
                    ),
                    mui.Box(
                        mui.IconButton(
                            mui.icon.MenuBook,
                            onClick=lambda: set_mode("manuel"),
                            sx={"color": "black"}
                        ),
                        mui.Typography(
                            "Manuel",
                            variant="h6",
                            color="black"
                        ),
                        display="flex", alignItems="center", sx={"marginRight": "10px"}
                    ),
                    mui.Box(
                        mui.IconButton(
                            mui.icon.List,
                            onClick=lambda: set_mode("dictionnaire"),
                            sx={"color": "black"}
                        ),
                        mui.Typography(
                            "Index",
                            variant="h6",
                            color="black"
                        ),
                        display="flex", alignItems="center",
                    ),
                    display="flex", alignItems="center", justifyContent="flex-end", width="600px"
                )

with st.spinner("Chargement..."):
    if st.session_state.mode == "pictortho":
        pictortho_page()

    elif st.session_state.mode == "tla":
        tla_page()

    elif st.session_state.mode == "manuel":
        st.markdown(st.session_state.manuel_file, unsafe_allow_html=True)

    elif st.session_state.mode == "dictionnaire":
        dictionnaire()

    for i in range(10):
        st.markdown("")

    st.markdown(
        f'<div style="text-align: center;">Tous pictogrammes présents sur Pictortho proviennent de la base de données d\'<a href="https://arasaac.org">ARASAAC</a> dont les termes d\'utilisations sont accessibles <a href="https://arasaac.org/terms-of-use">ici</a>.</div>',
        unsafe_allow_html=True)
    st.markdown(
        f'<div style="text-align: center;">Pictortho est disponible en libre accès sur <a href="https://github.com/victorC97/pictortho">Github</a>.</div>',
        unsafe_allow_html=True)
